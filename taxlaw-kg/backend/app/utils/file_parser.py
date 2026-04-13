"""
文档解析工具
支持 PDF、Word、Excel 格式
"""
import os
from typing import Optional
from pathlib import Path


class DocumentParser:
    """文档解析器"""

    SUPPORTED_TYPES = {
        "pdf": [".pdf"],
        "word": [".docx", ".doc"],
        "excel": [".xlsx", ".xls"],
        "text": [".txt", ".md", ".csv"]
    }

    def __init__(self, upload_dir: str = "./uploads"):
        self.upload_dir = upload_dir
        os.makedirs(upload_dir, exist_ok=True)

    def get_file_type(self, filename: str) -> Optional[str]:
        """根据文件扩展名判断文件类型"""
        ext = Path(filename).suffix.lower()
        for file_type, extensions in self.SUPPORTED_TYPES.items():
            if ext in extensions:
                return file_type
        return None

    def parse(self, file_path: str) -> str:
        """解析文档返回文本"""
        file_type = self.get_file_type(file_path)

        if file_type is None:
            raise ValueError(f"不支持的文件类型: {file_path}")

        if file_type == "pdf":
            return self._parse_pdf(file_path)
        elif file_type in ["docx", "doc"]:
            return self._parse_word(file_path)
        elif file_type in ["xlsx", "xls"]:
            return self._parse_excel(file_path)
        else:
            return self._parse_text(file_path)

    def _parse_pdf(self, file_path: str) -> str:
        """解析 PDF"""
        try:
            import pymupdf
            doc = pymupdf.open(file_path)
            text_parts = []
            for page in doc:
                text_parts.append(page.get_text())
            doc.close()
            return "\n".join(text_parts)
        except ImportError:
            # 备用方案：使用 pdfplumber
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    text_parts = []
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                    return "\n".join(text_parts)
            except ImportError:
                raise ImportError("请安装 pymupdf 或 pdfplumber 来解析 PDF 文件")

    def _parse_word(self, file_path: str) -> str:
        """解析 Word 文档"""
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = []
            for para in doc.paragraphs:
                paragraphs.append(para.text)

            # 提取表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    paragraphs.append(" | ".join(row_text))

            return "\n".join(paragraphs)
        except ImportError:
            raise ImportError("请安装 python-docx 来解析 Word 文档")

    def _parse_excel(self, file_path: str) -> str:
        """解析 Excel 文件"""
        try:
            import openpyxl
            text_parts = []

            wb = openpyxl.load_workbook(file_path, data_only=True)
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text_parts.append(f"[Sheet: {sheet_name}]")

                for row in sheet.iter_rows(values_only=True):
                    # 过滤 None 值并转为字符串
                    row_values = [str(v) if v is not None else "" for v in row]
                    if any(v for v in row_values):  # 只添加非空行
                        text_parts.append(" | ".join(row_values))

            wb.close()
            return "\n".join(text_parts)
        except ImportError:
            raise ImportError("请安装 openpyxl 来解析 Excel 文件")

    def _parse_text(self, file_path: str) -> str:
        """解析文本文件"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue

        raise ValueError(f"无法解码文件: {file_path}")


# 全局实例
_document_parser: Optional[DocumentParser] = None


def get_document_parser() -> DocumentParser:
    """获取文档解析器实例"""
    global _document_parser
    if _document_parser is None:
        from app.config import get_settings
        settings = get_settings()
        _document_parser = DocumentParser(upload_dir=settings.upload_dir)
    return _document_parser
